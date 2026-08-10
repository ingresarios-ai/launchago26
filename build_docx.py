import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def create_docx():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_sub = title_p.add_run("EL JUEGO MENTAL DEL DINERO · GUIÓN DE PRODUCCIÓN\n")
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Blue
    
    run_title = title_p.add_run("LIVE DÍA 1 — LA ANATOMÍA DEL SABOTEADOR")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark Navy

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(18)
    r_meta = subtitle_p.add_run("45 minutos · YouTube Live · Guión completo optimizado para máxima retención e interacción")
    r_meta.font.size = Pt(11)
    r_meta.font.italic = True
    r_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Read markdown content
    with open("Guion_Produccion_Live_Dia_1_El_Saboteador.md", "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_headers = []
    table_rows = []

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if not in_table or not table_headers:
            return
        
        table = doc.add_table(rows=len(table_rows) + 1, cols=len(table_headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, "E2E8F0")

        # Header
        hdr_cells = table.rows[0].cells
        for i, h_text in enumerate(table_headers):
            hdr_cells[i].text = h_text.strip()
            set_cell_background(hdr_cells[i], "0F172A") # Dark Navy
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Rows
        for r_idx, row_data in enumerate(table_rows):
            row_cells = table.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_value in enumerate(row_data):
                if c_idx < len(row_cells):
                    row_cells[c_idx].text = cell_value.strip()
                    set_cell_background(row_cells[c_idx], bg_color)
                    p = row_cells[c_idx].paragraphs[0]
                    for run in p.runs:
                        run.font.size = Pt(9.5)

        doc.add_paragraph() # spacing after table
        in_table = False
        table_headers = []
        table_rows = []

    for line in lines[4:]: # Skip top title lines
        sline = line.strip()

        # Handle Tables
        if sline.startswith("|"):
            parts = [p.strip() for p in sline.split("|")[1:-1]]
            if not parts or all(p.replace(":", "").replace("-", "") == "" for p in parts):
                continue # separator line
            if not in_table:
                in_table = True
                table_headers = parts
            else:
                table_rows.append(parts)
            continue
        else:
            if in_table:
                flush_table()

        if not sline:
            continue

        if sline.startswith("---"):
            continue

        # Headings
        if sline.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(sline.replace("## ", "").replace("📌 ", "").replace("🗺️ ", "").replace("📜 ", "").replace("🛠️ ", ""))
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(sline.replace("### ", "").replace("🔴 ", "").replace("⚡ ", "").replace("🎙️ ", ""))
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        elif sline.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            text = sline.replace("> ", "")
            # Callouts or quotes
            run = p.add_run(text.replace("**", ""))
            run.font.italic = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif sline.startswith("- ") or sline.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            content = sline[2:]
            # Bold processing
            parts = content.split("**")
            for idx, part in enumerate(parts):
                r = p.add_run(part)
                if idx % 2 == 1:
                    r.bold = True
        elif sline.startswith("1. ") or sline.startswith("2. ") or sline.startswith("3. ") or sline.startswith("4. ") or sline.startswith("5. ") or sline.startswith("6. "):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            content = sline[3:]
            parts = content.split("**")
            for idx, part in enumerate(parts):
                r = p.add_run(part)
                if idx % 2 == 1:
                    r.bold = True
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(4)
            
            parts = sline.split("**")
            for idx, part in enumerate(parts):
                r = p.add_run(part)
                if idx % 2 == 1:
                    r.bold = True

    if in_table:
        flush_table()

    output_filename = "Guion_Produccion_Live_Dia_1_El_Saboteador.docx"
    doc.save(output_filename)
    print(f"File saved successfully as {output_filename}")

if __name__ == "__main__":
    create_docx()
